from typing import Any, Dict, Optional
from docx import Document


def modify_docx(
    input_file_path: str,
    modifications: Dict[str, Any],
    output_file_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    Modifies a DOCX document with specified changes.

    Args:
        input_file_path (str): The path to the input DOCX file.
        modifications (Dict[str, Any]): The modifications to apply. Supported modifications include:
            - 'replace_text': Replaces occurrences of text. Requires 'old' and 'new' keys.
            - 'add_paragraph': Adds a new paragraph with the given text.
        output_file_path (str, optional): The path to save the modified DOCX file.

    Returns:
        Dict[str, Any]: A dictionary containing 'output', 'error', and 'error_code'.
    """
    try:
        # Load the document
        doc = Document(input_file_path)

        # Apply modifications (e.g., replace text, add paragraphs, etc.)
        for action, content in modifications.items():
            if action == 'replace_text':
                for paragraph in doc.paragraphs:
                    if content['old'] in paragraph.text:
                        paragraph.text = paragraph.text.replace(content['old'], content['new'])
            elif action == 'add_paragraph':
                doc.add_paragraph(content)

        if output_file_path:
            doc.save(output_file_path)
            return {
                "output": {
                    "message": f"Document modified and saved to '{output_file_path}'."
                },
                "error": None,
                "error_code": "SUCCESS"
            }
        else:
            return {
                "output": {
                    "message": "Modifications applied."
                },
                "error": None,
                "error_code": "SUCCESS"
            }

    except Exception as e:
        return {
            "output": None,
            "error": f"UNEXPECTED_ERROR: {e}",
            "error_code": "UNEXPECTED_ERROR"
        }